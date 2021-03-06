# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from cStringIO import StringIO
from django.http import HttpResponse
import os


def genWord(vals):
    #Windows
    path = '/home/rocco/www/intiSoft/Plantillas/'
    #path = '/home/nati/Desktop/intiSoft/Plantillas/'
    document = Document(path + vals['plantilla'])

    table = document.tables[0]

    # =======================================
    # ========== DATOS DEL ANEXO ============
    # =======================================

    # Area
    table.column_cells(0)[0].paragraphs[0].runs[0].add_text(vals['area'] or '')
    # Solicitante
    table.column_cells(1)[2].paragraphs[0].runs[0].add_text(vals['solicitante'] or '')
    # Contacto
    table.column_cells(4)[2].paragraphs[0].runs[0].add_text(vals['contacto'] or '')
    # Email
    table.column_cells(1)[3].paragraphs[0].runs[0].add_text(vals['email'] or '')

    if vals['plantilla'] != 'FO08–R02.docx':
        # Anexo
        table.column_cells(2)[0].paragraphs[0].runs[0].add_text(vals['codigo'] or '')
        # Presupuesto
        text = table.column_cells(0)[5].paragraphs[0].runs[0].text.split(',')
        text[0] = text[0].strip() + ' ' + vals['presupuesto']
        text = ','.join(text)
        table.column_cells(0)[5].paragraphs[0].runs[0].text = text

    if vals['plantilla'] in ['FO01–R06.docx', 'FO06–R02.docx']:
        # Fecha estimada de inicio
        table.column_cells(0)[6].paragraphs[1].runs[1].add_text(vals['fecha_inicio'])
        # Fecha estimada de fin
        table.column_cells(0)[6].paragraphs[12].runs[1].add_text(vals['fecha_fin'])
    elif vals['plantilla'] == 'FO07–R02.docx':
        # Fecha estimada de inicio
        table.column_cells(0)[6].paragraphs[8].runs[1].add_text(vals['fecha_inicio'])
        # Fecha estimada de fin
        table.column_cells(0)[6].paragraphs[16].runs[1].add_text(vals['fecha_fin'])
    elif vals['plantilla'] == 'FO08–R02.docx':
        # Anexo
        table.column_cells(0)[5].paragraphs[0].runs[0].add_text('\t' + vals['codigo'])
        # Presupuesto
        table.column_cells(0)[5].paragraphs[2].runs[0].add_text('\t' + vals['presupuesto'])
    elif vals['plantilla'] == 'FO09–R01.docx':
        # Fecha estimada de fin
        table.column_cells(0)[6].paragraphs[2].runs[1].add_text(vals['fecha_fin'])

    f = StringIO()
    document.save(f)
    f.seek(0)

    response = HttpResponse(f.getvalue(), content_type='text/docx')
    response['Content-Disposition'] = "attachment; filename=Anexo_%s.docx" % (vals['codigo'] or '')

    return response



def genSOT(vals):
    #Windows
    path = '/home/rocco/www/intiSoft/Plantillas/'

    document = Document(path + vals['plantilla'])

    style = document.styles['Normal']
    font = style.font
    font.name = 'Arial Narrow'
    font.size = Pt(8)

    # Documento temporal para la tabla de OTs
    tmp_doc = Document()
    ot_table = tmp_doc.add_table(0, 17)
    ot_table.alignment = 0
    table = document.tables[0]
    table_xml = document._part._element.body[0]
    # Necesario para ajustar el ancho de las columnas
    ot_table.autofit = False


    #=======================================
    #========== DATOS DE LA SOT ============
    #=======================================
    # Nro UT Ejecutora
    table.column_cells(3)[1].paragraphs[0].add_run(vals['nro_ejecutor'] or '')
    # UT Ejecutora
    table.column_cells(4)[1].paragraphs[0].add_run(vals['ejecutor'] or '')
    # Nro UT Deudora
    table.column_cells(10)[1].paragraphs[0].add_run(vals['nro_deudor'] or '')
    # UT Deudora
    table.column_cells(12)[1].paragraphs[0].add_run(vals['deudor'] or '')
    # Fecha de apertura
    table.column_cells(1)[3].paragraphs[0].add_run(vals['fecha_apertura'] or '')
    # OT Codigo
    table.column_cells(9)[3].paragraphs[0].add_run(vals['ot'] or '')
    # Expediente
    table.column_cells(9)[4].paragraphs[0].add_run(vals['expediente'] or '')
    # SOT Codigo
    table.column_cells(2)[5].paragraphs[0].add_run(vals['codigo'] or '')
    # Usuario final
    table.column_cells(5)[5].paragraphs[0].add_run(vals['usuario_final'] or '')
    table.column_cells(5)[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    ## OT's
    n = 1
    if vals['ofertatec']:
        cod, det, tipo_servicio, cantidad, precio, precio_total = vals['ofertatec'][0]
        table.column_cells(0)[8].paragraphs[0].add_run(cod + '\n' + det or '')
        table.column_cells(0)[8].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        table.column_cells(3)[8].paragraphs[0].add_run(tipo_servicio or '')
        table.column_cells(5)[8].paragraphs[0].add_run(str(cantidad) or '')
        table.column_cells(8)[8].paragraphs[0].add_run(str(precio) or '')
        table.column_cells(16)[8].paragraphs[0].add_run(str(precio_total) or '')
        for n, (cod, det, tipo_servicio, cantidad, precio, precio_total) in enumerate(vals['ofertatec'][1:], 2):
            # Creo y completo una fila en el documento temporal
            nrow = ot_table.add_row()
            # Ajusto el width
            for i in range(16):
                nrow.cells[i].width = table.column_cells(i)[7].width
            # Mergeo las celdas
            nrow.cells[0].merge(nrow.cells[2])
            nrow.cells[3].merge(nrow.cells[4])
            nrow.cells[5].merge(nrow.cells[7])
            nrow.cells[8].merge(nrow.cells[15])
            # Completo las columnas
            nrow.cells[0].paragraphs[0].add_run(cod + '\n' + det)
            nrow.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            nrow.cells[3].paragraphs[0].add_run(tipo_servicio or '')
            nrow.cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            nrow.cells[5].paragraphs[0].add_run(str(cantidad) or '')
            nrow.cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            nrow.cells[8].paragraphs[0].add_run(str(precio) or '')
            nrow.cells[8].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            nrow.cells[16].paragraphs[0].add_run(str(precio_total) or '')
            nrow.cells[16].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            ## Inserto la fila en el documento original
            table_xml.insert(n + 9, nrow._tr)

        table.column_cells(16)[n + 8].paragraphs[0].add_run(str(vals['arancel_previsto']) or '')
        table.column_cells(16)[n + 26].paragraphs[0].add_run(str(vals['importe_bruto']) or '')
        table.column_cells(16)[n + 27].paragraphs[0].add_run(str(vals['descuento']) or '')
        table.column_cells(16)[n + 28].paragraphs[0].add_run(str(vals['importe_neto']) or '')

    table.column_cells(1)[n + 8].paragraphs[0].add_run(vals['fecha_prevista'] or '')

    f = StringIO()
    document.save(f)
    f.seek(0)

    response = HttpResponse(f.getvalue(), content_type='text/docx')
    response['Content-Disposition'] = "attachment; filename=%s.docx" % ('SOT-' + vals['codigo'] or 'SOT')

    return response


def genRUT(vals):
    #Windows
    path = '/home/rocco/www/intiSoft/Plantillas/'

    document = Document(path + vals['plantilla'])

    style = document.styles['Normal']
    font = style.font
    font.name = 'Arial Narrow'
    font.size = Pt(8)

    # Documento temporal para la tabla de OTs
    tmp_doc = Document()
    ot_table = tmp_doc.add_table(0, 14)
    ot_table.alignment = 0

    table = document.tables[0]
    table_xml = document._part._element.body[0]
    # Necesario para ajustar el ancho de las columnas
    ot_table.autofit = False

    #=======================================
    #========== DATOS DE LA RUT ============
    #=======================================
    # Nro UT Ejecutora
    table.column_cells(1)[1].paragraphs[0].add_run(vals['nro_ejecutor'] or '')
    # UT Ejecutora
    table.column_cells(4)[1].paragraphs[0].add_run(vals['ejecutor'] or '')
    # Nro UT Deudora
    table.column_cells(7)[1].paragraphs[0].add_run(vals['nro_deudor'] or '')
    # UT Deudora
    table.column_cells(11)[1].paragraphs[0].add_run(vals['deudor'] or '')
    # Fecha de apertura
    table.column_cells(1)[3].paragraphs[0].add_run(vals['fecha_apertura'] or '')
    # RUT Codigo
    table.column_cells(1)[5].paragraphs[0].add_run(vals['codigo'] or '')
    ## OT's
    n = 1
    if vals['ofertatec']:
        cod, det, tipo_servicio, cantidad, precio, precio_total = vals['ofertatec'][0]
        table.column_cells(0)[8].paragraphs[0].add_run(cod + '\n' + det or '')
        table.column_cells(0)[8].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        table.column_cells(3)[8].paragraphs[0].add_run(tipo_servicio or '')
        table.column_cells(5)[8].paragraphs[0].add_run(str(cantidad) or '')
        table.column_cells(6)[8].paragraphs[0].add_run(str(precio) or '')
        table.column_cells(9)[8].paragraphs[0].add_run(str(precio_total) or '')
        for n, (cod, det, tipo_servicio, cantidad, precio, precio_total) in enumerate(vals['ofertatec'][1:], 2):
            # Creo y completo una fila en el documento temporal
            nrow = ot_table.add_row()
            # Ajusto el width
            for i in range(14):
                nrow.cells[i].width = table.column_cells(i)[7].width
            # Mergeo las celdas
            nrow.cells[0].merge(nrow.cells[2])
            nrow.cells[3].merge(nrow.cells[4])
            nrow.cells[6].merge(nrow.cells[8])
            nrow.cells[9].merge(nrow.cells[13])
            # Completo las columnas
            nrow.cells[0].paragraphs[0].add_run(cod + '\n' + det)
            nrow.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            nrow.cells[3].paragraphs[0].add_run(tipo_servicio or '')
            nrow.cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            nrow.cells[5].paragraphs[0].add_run(str(cantidad) or '')
            nrow.cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            nrow.cells[6].paragraphs[0].add_run(str(precio) or '')
            nrow.cells[6].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            nrow.cells[9].paragraphs[0].add_run(str(precio_total) or '')
            nrow.cells[9].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            ## Inserto la fila en el documento original
            table_xml.insert(n + 9, nrow._tr)

        table.column_cells(13)[n + 8].paragraphs[0].add_run(str(vals['arancel_previsto']) or '')
        table.column_cells(13)[n + 26].paragraphs[0].add_run(str(vals['importe_bruto']) or '')
        table.column_cells(13)[n + 27].paragraphs[0].add_run(str(vals['descuento']) or '')
        table.column_cells(13)[n + 28].paragraphs[0].add_run(str(vals['importe_neto']) or '')

    table.column_cells(1)[n + 8].paragraphs[0].add_run(vals['fecha_prevista'] or '')

    f = StringIO()
    document.save(f)
    f.seek(0)

    response = HttpResponse(f.getvalue(), content_type='text/docx')
    response['Content-Disposition'] = "attachment; filename=%s.docx" % ('RUT-' + vals['codigo'] or 'SOT')

    return response


def genSI(vals):
    #Windows
    path = '/home/rocco/www/intiSoft/Plantillas/'

    document = Document(path + vals['plantilla'])

    style = document.styles['Normal']
    font = style.font
    font.name = 'Arial Narrow'
    font.size = Pt(8)

    # Documento temporal para la tabla de OTs
    tmp_doc = Document()
    ot_table = tmp_doc.add_table(0, 28)
    ot_table.alignment = 0

    table = document.tables[0]
    table_xml = document._part._element.body[0]
    # Necesario para ajustar el ancho de las columnas
    ot_table.autofit = False

    #=======================================
    #========== DATOS DE LA SI ============
    #=======================================
    # UT Ejecutora
    table.column_cells(1)[2].paragraphs[0].add_run(vals['ejecutor'] or '')
    # UT Deudora
    table.column_cells(21)[2].paragraphs[0].add_run(vals['solicitante'] or '')
    # Fecha de apertura
    table.column_cells(1)[4].paragraphs[0].add_run(vals['fecha_apertura'] or '')
    # RUT Codigo
    table.column_cells(1)[6].paragraphs[0].add_run(vals['codigo'] or '')
    ## OT's
    n = 1
    if vals['ofertatec']:
        cod, det, tipo_servicio, cantidad, precio, precio_total = vals['ofertatec'][0]
        table.column_cells(0)[9].paragraphs[0].add_run(cod + '\n' + det or '')
        table.column_cells(0)[9].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        table.column_cells(7)[9].paragraphs[0].add_run(tipo_servicio or '')
        table.column_cells(26)[9].paragraphs[0].add_run(str(cantidad) or '')
        for n, (cod, det, tipo_servicio, cantidad, precio, precio_total) in enumerate(vals['ofertatec'][1:], 2):
            # Creo y completo una fila en el documento temporal
            nrow = ot_table.add_row()
            # Ajusto el width
            for i in range(27):
                nrow.cells[i].width = table.column_cells(i)[8].width
            # Mergeo las celdas
            nrow.cells[0].merge(nrow.cells[6])
            nrow.cells[7].merge(nrow.cells[25])
            nrow.cells[26].merge(nrow.cells[27])
            # Completo las columnas
            nrow.cells[0].paragraphs[0].add_run(cod + '\n' + det)
            nrow.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            nrow.cells[7].paragraphs[0].add_run(tipo_servicio or '')
            nrow.cells[7].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            nrow.cells[26].paragraphs[0].add_run(str(cantidad) or '')
            nrow.cells[26].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Inserto la fila en el documento original
            table_xml.insert(n + 10, nrow._tr)

    if vals.has_key('fecha_inicio'):
        table.column_cells(1)[n + 9].paragraphs[0].add_run(vals['fecha_inicio'] or '')
        table.column_cells(1)[n + 10].paragraphs[0].add_run(vals['fecha_fin'] or '')

    for m, (tarea, horas) in enumerate(vals['tarea']):
        table.column_cells(0)[n + m + 18].paragraphs[0].add_run(tarea)
        table.column_cells(0)[n + m + 18].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.column_cells(20)[n + m + 18].paragraphs[0].add_run(str(horas))
        table.column_cells(20)[n + m + 18].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    f = StringIO()
    document.save(f)
    f.seek(0)

    response = HttpResponse(f.getvalue(), content_type='text/docx')
    response['Content-Disposition'] = "attachment; filename=%s.docx" % ('SI-' + vals['codigo'] or 'SI')

    return response
